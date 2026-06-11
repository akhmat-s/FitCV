from __future__ import annotations

import os
import shutil
import zipfile
from io import BytesIO
from pathlib import Path

import docx  # python-docx
import fitz  # PyMuPDF
from docx.opc.exceptions import PackageNotFoundError

from schemas import ACCEPTED_FORMATS


class UnsupportedFormatError(ValueError):
    """Raised when an uploaded file's extension is not in ``ACCEPTED_FORMATS``.

    A domain exception so ``extract.build_extract`` can catch the format-guard
    failure specifically and translate it to an ``ErrorResponse`` at
    ``stage=parse``, without catching unrelated errors.
    """


class CorruptFileError(ValueError):
    """Raised when an accepted-format file cannot be parsed (corrupt/unreadable).

    The underlying parsers raise library-specific errors
    (``fitz.FileDataError``/``fitz.EmptyFileError`` for PDFs, ``zipfile.BadZipFile``
    for DOCX). We wrap them in one domain exception so ``extract.build_extract`` can
    translate a bad file to an ``ErrorResponse`` at ``stage=parse`` without catching
    unrelated provider errors (no bare ``except Exception``).
    """


def extract_text(data: bytes, filename: str) -> str:
    """Extract plain text from CV bytes entirely in memory (no disk I/O).

    Stateless parse path: bytes in, text out, nothing written.
    `filename` is used only to route on the extension.

    The legacy folder-batch ``PDFParser`` below is kept intact for the existing
    offline workflow.

    Args:
        data: Raw uploaded file bytes.
        filename: Original filename, used to detect the extension.

    Returns:
        The extracted text.

    Raises:
        UnsupportedFormatError: When the extension is not in ``ACCEPTED_FORMATS``.
        CorruptFileError: When an accepted-format file cannot be parsed.
    """
    extension = Path(filename).suffix.lower().lstrip(".")

    if extension not in ACCEPTED_FORMATS:
        accepted = ", ".join(ACCEPTED_FORMATS).upper()
        raise UnsupportedFormatError(
            f"Unsupported file format '.{extension or '?'}'. Upload one of: {accepted}."
        )

    if extension == "pdf":
        return _extract_pdf_text(data)
    if extension == "docx":
        return _extract_docx_text(data)
    return _extract_txt_text(data)


def _extract_pdf_text(data: bytes) -> str:
    """Read all page text from in-memory PDF bytes via PyMuPDF's stream API.

    Raises:
        CorruptFileError: When the bytes are not a readable PDF stream.
    """
    try:
        pages: list[str] = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc.pages():
                pages.append(page.get_text("text"))
        return "".join(pages)
    except (fitz.FileDataError, fitz.EmptyFileError) as exc:
        raise CorruptFileError("Couldn't read this PDF; the file may be corrupt.") from exc


def _extract_docx_text(data: bytes) -> str:
    """Read paragraph + table text from in-memory DOCX bytes via python-docx.

    Raises:
        CorruptFileError: When the bytes are not a readable DOCX (zip) container.
    """
    # A non-zip is BadZipFile, but a valid ZIP that is not an OOXML package raises
    # PackageNotFoundError, and one missing an internal part raises KeyError. All
    # three mean "unreadable DOCX" and must halt at stage=parse, not crash.
    try:
        document = docx.Document(BytesIO(data))
    except (zipfile.BadZipFile, PackageNotFoundError, KeyError) as exc:
        raise CorruptFileError("Couldn't read this DOCX; the file may be corrupt.") from exc
    parts: list[str] = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt_text(data: bytes) -> str:
    """Decode plain-text CV bytes as UTF-8, replacing undecodable bytes."""
    return data.decode("utf-8", errors="replace")


class PDFParser:
    def __init__(
        self,
        source_folder: str = 'data/cv_example',
        result_folder: str = 'data/cv_txt',
    ):
        self.source_folder = Path(source_folder)
        self.result_folder = Path(os.getcwd(), result_folder)

        self.data_path = self.result_folder
        self.data_path.mkdir(parents=True, exist_ok=True)

    def parse(self):
        """Parses PDF and DOCX files, extracts text, saving document structure.
        All pages from one document are saved to a single text file.
        For TXT files, they are just copied to the destination folder.

        Note:
            If a file cannot be opened, outputs a message with the filename.
        """
        for file_path in self.source_folder.iterdir():
            file_extension = file_path.suffix.lower()

            # Output filename - same as input but with .txt extension
            output_filename = file_path.stem + '.txt'
            text_filepath = self.data_path / output_filename

            try:
                if file_extension == '.pdf':
                    self._process_pdf(file_path, text_filepath)
                elif file_extension == '.docx':
                    self._process_docx(file_path, text_filepath)
                elif file_extension == '.txt':
                    self._process_txt(file_path, text_filepath)
                else:
                    continue  # Skip other file types

                print(f'File processed: {file_path.name} -> {output_filename}')

            except Exception as e:
                print(f'Error processing file {file_path.name}: {str(e)}')

    def _process_pdf(self, file_path, text_filepath):
        """Process a PDF file and extract its text."""
        all_pages_text = []

        with fitz.open(file_path) as doc:
            for page in doc.pages():
                # Get page text in a format that preserves structure
                page_text = page.get_text("text")
                all_pages_text.append(page_text)

        # Combine text from all pages and save to a single file
        full_text = ''.join(all_pages_text)
        text_filepath.write_text(full_text, encoding='utf-8')

    def _process_docx(self, file_path, text_filepath):
        """Process a DOCX file and extract its text."""
        doc = docx.Document(file_path)

        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs]

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)

        # Join all text with newlines and save
        full_text = '\n'.join(paragraphs)
        text_filepath.write_text(full_text, encoding='utf-8')

    def _process_txt(self, file_path, text_filepath):
        """Copy TXT file to the destination folder."""
        shutil.copy2(file_path, text_filepath)


if __name__ == '__main__':
    parser = PDFParser()
    parser.parse()
